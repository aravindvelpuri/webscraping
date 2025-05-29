from docx import Document
from docx.shared import Inches
from datetime import datetime

def generate_word_report(df):
    doc = Document()
    doc.add_heading('Andhra Pradesh News Report', 0)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'Source'
    hdr_cells[2].text = 'Severity'
    hdr_cells[3].text = 'URL'

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['title']
        cells[1].text = row['source']
        cells[2].text = row['severity']
        cells[3].text = row['url']

    filename = f"andhra_news_report_{datetime.now().strftime('%Y-%m-%d')}.docx"
    doc.save(filename)
    return filename
